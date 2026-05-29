from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX = OUT_DIR / "IronBrother_Contract_Security_Audit_Report_New_2026-05-20.docx"

FONT = "Microsoft YaHei"
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(96, 96, 96)
GREEN = RGBColor(0, 97, 0)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "E2F0D9"
GRID = "B7C4D6"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size=None, color=None, bold=None):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tcw = tc_pr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tc_pr.append(tcw)
            tcw.set(qn("w:w"), str(int(width * 1440)))
            tcw.set(qn("w:type"), "dxa")


def add_para(doc, text="", size=10.8, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r, size=size, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 12.5, bold=True, color=BLUE if level == 1 else DARK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=10.5)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, widths)
    table_borders(table)
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], header_fill)
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=9.2)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, body, fill=LIGHT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [6.35])
    table_borders(table, color="C6E0B4", size="5")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, 11, True, GREEN)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body)
    set_run_font(r2, 10.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
    set_style_font(doc.styles[style_name], 11)
doc.styles["Normal"].paragraph_format.space_after = Pt(6)
doc.styles["Normal"].paragraph_format.line_spacing = 1.12

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run_font(header.add_run("IronBrother 合约安全审计报告 | 新版"), 9, color=MUTED)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(footer.add_run("BSC Mainnet Contract Review - 2026-05-20"), 8.5, color=MUTED)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
set_run_font(title.add_run("IronBrother 合约安全审计报告"), 24, True, RGBColor(0, 0, 0))
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
set_run_font(subtitle.add_run("新版 / BSC Mainnet 已部署验收 / 合约专项安全结论"), 13, color=MUTED)

add_table(
    doc,
    ["项目", "内容"],
    [
        ("报告日期", "2026-05-20 / Asia/Shanghai"),
        ("审计对象", "contracts/IronBrother.sol；BSC Mainnet UUPS Proxy"),
        ("主网 Proxy", "0x25cCc567C5a72Bb164d0e75969ff6141a7d735E3"),
        ("当前 Implementation", "0xe6629Ed50d11F921C0717Cb2d9B36C9231e0f7B8"),
        ("Proxy 升级交易", "0x812da3b1b194ed4cd9ec9322966095b75cffa4baead8792c360c3cacc2445a48"),
        ("Implementation 部署交易", "0xc5a314f4a8c4625ad270c20f035d71e75b75d7e39fce4a9088bcd726ab2caa00"),
    ],
    [1.55, 4.95],
    header_fill=LIGHT_GRAY,
    font_size=8.8,
)

add_callout(
    doc,
    "审计结论",
    "本次新版合约审计确认：已发现的推荐关系成环风险与动态奖励关闭日异常分支均已完成代码修复、测试验证和 BSC Mainnet 部署验收。当前主网 proxy 指向新 implementation；链上 future day 调用正确拒绝并返回 `day not closed`，合约专项测试通过。",
)

add_heading(doc, "1. 审计范围与修复事项")
add_table(
    doc,
    ["编号", "原风险", "修复状态", "复审结论"],
    [
        ("R-01", "推荐关系可以形成环，影响动态奖励完整性", "已修复", "_bindReferrer 绑定前沿上级链检查，发现回到当前用户即 revert `referrer cycle`。"),
        ("R-02", "动态奖励关闭日判断存在未来日异常分支", "已修复", "_isDynamicRewardDayClosed 仅返回 `day < currentLocalDay()`，未来 day 不再被视为已关闭。"),
    ],
    [0.55, 2.1, 0.9, 2.95],
    font_size=8.8,
)

add_heading(doc, "2. 主网部署与链上验收")
add_table(
    doc,
    ["检查项", "结果", "说明"],
    [
        ("Proxy implementation", "通过", "ERC1967 slot 当前为 0xe6629Ed50d11F921C0717Cb2d9B36C9231e0f7B8。"),
        ("字节码一致性", "通过", "链上 runtime 与本地 artifact 均为 24,524 bytes；忽略 UUPS immutable self 后哈希一致。"),
        ("未来 day 结算", "通过", "链上 staticCall 使用 futureDay=21593，futureDayRejectedWithDayNotClosed=true，正确 revert `day not closed`。"),
        ("合约专项测试", "通过", "npx hardhat test test\\IronBrother.test.cjs test\\dynamicRewards.test.cjs：25 passing。"),
        ("部署记录", "通过", "deployments/bsc.json 与 .openzeppelin/bsc.json 已同步到新 implementation。"),
    ],
    [1.45, 0.8, 4.25],
    header_fill=LIGHT_BLUE,
    font_size=8.8,
)

add_heading(doc, "3. 新版安全判断")
add_bullet(doc, "普通用户不能通过推荐关系构造环形上级链来重复放大动态奖励路径。")
add_bullet(doc, "动态奖励结算入口不再接受未来 day；`settleDynamicRewardForUser`、批量结算和 Manager bot 路径均受同一关闭日判断约束。")
add_bullet(doc, "本次修复没有新增存储变量，保持 UUPS 升级存储布局安全。")
add_bullet(doc, "合约运行时代码体积为 24,524 bytes，低于 24,576 bytes 主网限制。")
add_bullet(doc, "本次主网升级使用具备 DEFAULT_ADMIN_ROLE 的部署地址完成，升级后 proxy 指向新 implementation。")

add_heading(doc, "4. 关键源码证据")
add_table(
    doc,
    ["位置", "逻辑", "安全意义"],
    [
        ("contracts/IronBrother.sol:892", "_bindReferrer 中检查 `referrer != user`", "保留直接自推荐阻断。"),
        ("contracts/IronBrother.sol:894", "沿 `users[referrer].referrer` 向上遍历", "阻止间接推荐环。"),
        ("contracts/IronBrother.sol:895", "发现 `cursor == user` 即 revert `referrer cycle`", "防止 A->B->...->A。"),
        ("contracts/IronBrother.sol:979", "_isDynamicRewardDayClosed(day)", "统一动态奖励关闭日判断。"),
        ("contracts/IronBrother.sol:981", "返回 `day < currentDay`", "只允许已过去本地日结算。"),
    ],
    [1.75, 2.35, 2.4],
    font_size=8.5,
)

add_heading(doc, "5. 验证命令")
add_table(
    doc,
    ["命令 / 检查", "结果"],
    [
        ("npx hardhat compile", "通过，无代码体积超限警告。"),
        ("npx hardhat test test\\IronBrother.test.cjs test\\dynamicRewards.test.cjs", "通过，25 passing。"),
        ("链上 implementation slot 读取", "通过，当前 implementation 为 0xe6629Ed50d11F921C0717Cb2d9B36C9231e0f7B8。"),
        ("链上 future day staticCall", "通过，futureDayRejectedWithDayNotClosed=true，正确 revert `day not closed`。"),
    ],
    [3.1, 3.4],
    header_fill=LIGHT_GRAY,
    font_size=8.8,
)

add_heading(doc, "6. 运行建议")
add_para(doc, "本次新版审计结论为正向通过。后续仍建议将高权限操作纳入多签或时间锁，并持续监控奖励池余额、提现状态和动态奖励结算任务执行情况。这些属于运营治理建议，不影响本次两个合约逻辑问题的修复结论。")

doc.save(DOCX)
print(DOCX)
