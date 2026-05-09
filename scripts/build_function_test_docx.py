from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "功能测试文档.md"
OUTPUT = ROOT / "docs" / "功能测试文档.docx"

FONT = "Microsoft YaHei"
ACCENT = "1F4E79"
ACCENT_DARK = "173B5F"
LIGHT_BLUE = "EAF3FA"
LIGHT_GRAY = "F5F7FA"
LIGHT_WARN = "FFF3CD"
WHITE = "FFFFFF"
GRID = "CBD5E1"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_style(paragraph, size: float = 10.5, color: str | None = None):
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.12
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = GRID):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_padding(cell, top: int = 90, left: int = 110, bottom: int = 90, right: int = 110):
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = tc_pr.first_child_found_in("w:tcMar")
    if margin is None:
        margin = OxmlElement("w:tcMar")
        tc_pr.append(margin)
    for name, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = margin.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margin.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, inches: float):
    width = int(inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(inches)


def set_fixed_table_layout(table, widths: list[float]):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    old_grid = table._tbl.tblGrid
    if old_grid is not None:
        table._tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    table._tbl.insert(0, grid)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.2, color: str = "111827", align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    add_inline_runs(paragraph, text, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_padding(cell)
    set_cell_borders(cell)


def add_inline_runs(paragraph, text: str, size: float = 10.5, bold: bool = False, color: str | None = None):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        is_code = part.startswith("`") and part.endswith("`")
        value = part[1:-1] if is_code else part
        run = paragraph.add_run(value)
        set_run_font(run, size=size, bold=bold, color=color)
        if is_code:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.color.rgb = RGBColor.from_string("0F766E")


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        row = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in row):
            rows.append(row)
        i += 1
    return rows, i


def choose_widths(headers: list[str], col_count: int) -> list[float]:
    header_text = "|".join(headers)
    if col_count == 2:
        return [2.0, 8.1]
    if col_count == 3:
        return [1.5, 4.2, 4.4]
    if "用例 ID" in header_text:
        return [1.1, 2.0, 3.35, 3.7]
    if "业务功能" in header_text:
        return [1.35, 2.25, 3.6, 2.95]
    if "编号" in header_text and "差异点" in header_text:
        return [1.0, 2.6, 3.65, 2.9]
    return [10.15 / col_count] * col_count


def add_table(document: Document, rows: list[list[str]]):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    widths = choose_widths(rows[0], col_count)
    table = document.add_table(rows=0, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    table.style = "Table Grid"
    set_fixed_table_layout(table, widths)

    for row_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        for col_idx in range(col_count):
            text = row_data[col_idx] if col_idx < len(row_data) else ""
            cell = cells[col_idx]
            if col_idx < len(widths):
                set_cell_width(cell, widths[col_idx])
            if row_idx == 0:
                shade_cell(cell, ACCENT_DARK)
                set_cell_text(cell, text, bold=True, size=8.4, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                if row_idx % 2 == 0:
                    shade_cell(cell, LIGHT_GRAY)
                elif "GAP-" in row_data[0] or "P0" in text:
                    shade_cell(cell, LIGHT_WARN)
                set_cell_text(cell, text, size=8.0, align=WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    document.add_paragraph()


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("IronBrother 功能测试文档  |  第 ")
    set_run_font(run, 8.5, color="64748B")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    run = footer.add_run(" 页")
    set_run_font(run, 8.5, color="64748B")


def add_cover(document: Document):
    for _ in range(5):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("IronBrother 功能测试文档")
    set_run_font(run, 26, bold=True, color=ACCENT_DARK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("业务功能与实现逻辑对照 · 测试用例手册")
    set_run_font(run, 14, color=ACCENT)

    document.add_paragraph()
    meta = document.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.allow_autofit = False
    set_fixed_table_layout(meta, [1.4, 6.4])
    data = [
        ("版本", "v1.0"),
        ("日期", "2026-05-07"),
        ("测试对象", "IronBrother DApp / Admin 后台 / BSC Testnet 合约"),
        ("业务来源", "铁哥技术.xmind、README.md、IronBrother.sol、App.tsx"),
    ]
    for i, (key, value) in enumerate(data):
        cells = meta.rows[i].cells
        set_cell_width(cells[0], 1.4)
        set_cell_width(cells[1], 6.4)
        shade_cell(cells[0], LIGHT_BLUE)
        set_cell_text(cells[0], key, bold=True, size=9.6, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], value, size=9.6)

    document.add_page_break()


def add_heading(document: Document, text: str, level: int):
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {min(level, 3)}"
    paragraph.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, 18, bold=True, color=ACCENT_DARK)
    elif level == 2:
        set_run_font(run, 13.5, bold=True, color=ACCENT)
    else:
        set_run_font(run, 11.5, bold=True, color="334155")


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    add_footer(section)

    styles = document.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(10.5)

    add_cover(document)

    i = 0
    skip_intro_meta = True
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if skip_intro_meta and (stripped.startswith("版本：") or stripped.startswith("日期：") or stripped.startswith("测试对象：") or stripped.startswith("业务来源：")):
            i += 1
            continue
        if stripped.startswith("# "):
            i += 1
            continue
        if stripped.startswith("## "):
            skip_intro_meta = False
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(document, rows)
            continue
        if stripped.startswith("### "):
            add_heading(document, stripped[4:].strip(), 3)
        elif stripped.startswith("## "):
            add_heading(document, stripped[3:].strip(), 2)
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, stripped[2:].strip(), size=10.1)
            set_paragraph_style(paragraph, 10.1)
        elif re.match(r"^\d+\.\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            add_inline_runs(paragraph, re.sub(r"^\d+\.\s+", "", stripped), size=10.1)
            set_paragraph_style(paragraph, 10.1)
        else:
            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, stripped, size=10.5)
            set_paragraph_style(paragraph, 10.5)
        i += 1

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
